# ğŸ“¦ Install necessary packages before running (if not already installed)
# !pip install transformers streamlit keybert sentence-transformers PyMuPDF langdetect python-docx markdown

import streamlit as st
from transformers import pipeline, MBartForConditionalGeneration, MBart50TokenizerFast
from keybert import KeyBERT
from sentence_transformers import SentenceTransformer
import fitz  # PyMuPDF
from langdetect import detect
import re
from docx import Document
import markdown

# Initialize models
summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
keyword_model = KeyBERT(model=SentenceTransformer('all-MiniLM-L6-v2'))

# Load MBART for multilingual support
mbart_model = MBartForConditionalGeneration.from_pretrained("facebook/mbart-large-50-many-to-many-mmt")
mbart_tokenizer = MBart50TokenizerFast.from_pretrained("facebook/mbart-large-50-many-to-many-mmt")

# ğŸ“Œ Function to extract text from PDF
def extract_text_from_pdf(pdf_file):
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text

# ğŸ“Œ Function to detect language
def detect_language(text):
    return detect(text)

# ğŸ“Œ Summarize text (with safe token limit handling)
def generate_summary(text):
    if len(text.split()) < 30:
        return "Text too short to summarize."

    max_input_length = 1024  # BART max token limit (approx as words)
    words = text.split()
    if len(words) > max_input_length:
        text = ' '.join(words[:max_input_length])

    summary = summarizer(text, max_length=130, min_length=30, do_sample=False)
    return summary[0]['summary_text']

# ğŸ“Œ Keyword extraction
def extract_keywords(text):
    keywords = keyword_model.extract_keywords(text, keyphrase_ngram_range=(1, 2), stop_words='english', top_n=5)
    return [kw[0] for kw in keywords]

# ğŸ“Œ AI-powered Notes Generator (with safe token limit handling)
def generate_notes(text):
    max_input_length = 1024
    words = text.split()
    if len(words) > max_input_length:
        text_for_summary = ' '.join(words[:max_input_length])
    else:
        text_for_summary = text

    key_points = summarizer(text_for_summary, max_length=150, min_length=50, do_sample=False)[0]['summary_text']
    keywords = extract_keywords(text)
    action_items = re.findall(r'(?i)(?:should|must|need to|required to|responsible for)\s.*?\.', text)
    open_questions = re.findall(r'[^.?!]*\?', text)
    stats = re.findall(r'\b\d{1,3}%|\d+\.\d+|\d{1,3}\b', text)

    notes = {
        "Key Points": key_points,
        "Keywords": keywords,
        "Action Items": action_items,
        "Open Questions": open_questions,
        "Important Stats": stats
    }
    return notes

# ğŸ“Œ Export notes to DOCX
def export_notes_to_docx(notes):
    doc = Document()
    doc.add_heading("AI-Powered Notes", 0)
    for section, content in notes.items():
        doc.add_heading(section, level=1)
        if isinstance(content, list):
            for item in content:
                doc.add_paragraph(item, style='List Bullet')
        else:
            doc.add_paragraph(content)
    doc_path = "notes_output.docx"
    doc.save(doc_path)
    return doc_path

# ğŸ“Œ Export notes to Markdown
def export_notes_to_markdown(notes):
    md_content = "# AI-Powered Notes\n"
    for section, content in notes.items():
        md_content += f"\n## {section}\n"
        if isinstance(content, list):
            for item in content:
                md_content += f"- {item}\n"
        else:
            md_content += f"{content}\n"
    md_path = "notes_output.md"
    with open(md_path, "w", encoding='utf-8') as f:
        f.write(md_content)
    return md_path

# ğŸ“Œ Streamlit UI
st.title("ğŸ“‘ AI-Powered Text Summarizer & Notes Generator")
st.write("Upload a text or PDF file and get summaries, keywords, and structured notes.")

uploaded_file = st.file_uploader("Upload a PDF file", type="pdf")
text_input = st.text_area("Or paste your text here:")

if uploaded_file is not None:
    input_text = extract_text_from_pdf(uploaded_file)
elif text_input:
    input_text = text_input
else:
    input_text = ""

if st.button("Generate Insights") and input_text:
    lang = detect_language(input_text)
    st.write(f"Detected Language: {lang}")

    summary = generate_summary(input_text)
    st.subheader("ğŸ“„ Summary")
    st.write(summary)

    keywords = extract_keywords(input_text)
    st.subheader("ğŸ”‘ Keywords")
    st.write(", ".join(keywords))

    notes = generate_notes(input_text)
    st.subheader("ğŸ“ AI-Powered Notes")
    for section, content in notes.items():
        st.write(f"**{section}:**")
        st.write(content if content else "- None found.")

    docx_path = export_notes_to_docx(notes)
    md_path = export_notes_to_markdown(notes)

    with open(docx_path, "rb") as f:
        st.download_button(label="ğŸ“¥ Download Notes as DOCX", data=f, file_name="notes_output.docx")

    with open(md_path, "rb") as f:
        st.download_button(label="ğŸ“¥ Download Notes as Markdown", data=f, file_name="notes_output.md")

# ğŸ“Œ Future Scope Comments
st.sidebar.header("ğŸš€ Future Enhancements")
st.sidebar.write("- Integrate real-time Speech-to-Text transcription.")
st.sidebar.write("- Enhance Action Item detection with a fine-tuned classifier.")
